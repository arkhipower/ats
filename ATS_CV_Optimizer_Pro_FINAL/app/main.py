import os
import sys
from pathlib import Path
PROMPT_DIR = Path(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "prompts")))
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import os
import pdfplumber
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document as DocxDocument

# === Globals ===
cv_text = ""
jd_text = ""
output_text = None
audit_mode = None

# === Utils ===
def extract_text(file_path):
    if file_path.endswith('.pdf'):
        with pdfplumber.open(file_path) as pdf:
            return "\n".join([page.extract_text() or '' for page in pdf.pages])
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        return ""

def update_output(text):
    output_text.delete(1.0, tk.END)
    output_text.insert(tk.END, text)

# === Functions ===
def upload_cv():
    global cv_text
    file_path = filedialog.askopenfilename(filetypes=[("Documents", "*.pdf *.docx *.txt")])
    if file_path:
        cv_text = extract_text(file_path)
        update_output("✅ CV uploaded successfully.\n")

def upload_jd():
    global jd_text
    file_path = filedialog.askopenfilename(filetypes=[("Documents", "*.pdf *.docx *.txt")])
    if file_path:
        jd_text = extract_text(file_path)
        update_output("✅ Job Description uploaded successfully.\n")


from ai.ollama_client import run_ollama_prompt, set_model, run_prompt_from_file
from config.license_check import check_license_key

def run_ai_audit():
    global cv_text, jd_text
    if not cv_text or not jd_text:
        update_output("⚠️ Please upload both CV and JD.")
        return

    mode = audit_mode.get()
    prompt_file = {
        "consultant": PROMPT_DIR / "prompt_audit_consultant.txt",
        "client": PROMPT_DIR / "prompt_audit_client.txt",
        "self": PROMPT_DIR / "prompt_audit_self_audit.txt",
        "compare": PROMPT_DIR / "prompt_compare_cv_jd.txt",
        "optimize": "prompts/prompt_optimization.txt"
    }.get(mode)

    if not prompt_file:
        update_output("❌ Invalid audit mode.")
        return

    try:
        result = run_prompt_from_file(prompt_file, {
            "cv_text": cv_text,
            "jd_text": jd_text
        })
        update_output(result)
    except Exception as e:
        update_output(f"[ERROR] {str(e)}")

def generate_white_keywords():
    if not cv_text:
        update_output("⚠️ Please upload CV first.")
        return
    industry_keywords = ["HSE", "Oil & Gas", "ISO 45001", "Permit to Work", "Risk Assessment"]
    white_text = "\n\n[White Keywords (Invisible to reader)]\n" + " ".join(industry_keywords)
    update_output(cv_text + white_text)

def generate_cover_letter():
    if not cv_text or not jd_text:
        update_output("⚠️ Please upload both CV and JD.")
        return
    prompt_path = PROMPT_DIR / "prompt_cover_letter.txt"
    result = run_prompt_from_file(prompt_path, {
        "cv_text": cv_text,
        "jd_text": jd_text
    })
    update_output(result)
def run_scoring():
    if not cv_text or not jd_text:
        update_output("⚠️ Please upload both CV and JD.")
        return
    docs = [cv_text, jd_text]
    vectorizer = TfidfVectorizer()
    tfidf = vectorizer.fit_transform(docs)
    score = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
    update_output(f"📊 CV vs JD TF-IDF Similarity Score: {score:.2f}")

def export_docx():
    if not output_text.get(1.0, tk.END).strip():
        messagebox.showinfo("Export", "No content to export.")
        return
    file_path = filedialog.asksaveasfilename(defaultextension=".docx")
    if file_path:
        doc = DocxDocument()
        doc.add_paragraph(output_text.get(1.0, tk.END))
        doc.save(file_path)
        messagebox.showinfo("Export", f"Saved to {file_path}")

def save_result():
    if not output_text.get(1.0, tk.END).strip():
        messagebox.showinfo("Save", "No content to save.")
        return
    file_path = filedialog.asksaveasfilename(defaultextension=".txt")
    if file_path:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(output_text.get(1.0, tk.END))
        messagebox.showinfo("Save", f"Saved to {file_path}")

def clear_fields():
    global cv_text, jd_text
    cv_text, jd_text = "", ""
    update_output("")


def check_license():
    key = tk.simpledialog.askstring("License Check", "Enter license key:")
    if not key:
        return
    if check_license_key(key):
        messagebox.showinfo("License", "✅ License key valid.")
    else:
        messagebox.showerror("License", "❌ Invalid license key.")

    messagebox.showinfo("License", "License valid. (Simulation)")

# === UI ===

def generate_cover_letter():
    if not cv_text or not jd_text:
        update_output("⚠️ Please upload both CV and JD.")
        return
    prompt_path = PROMPT_DIR / "prompt_cover_letter.txt"
    result = run_prompt_from_file(prompt_path, {
        "cv_text": cv_text,
        "jd_text": jd_text
    })
    update_output(result)
def run_optimization():
    global cv_text, jd_text
    if not cv_text or not jd_text:
        update_output("⚠️ Please upload both CV and JD.")
        return
    result = run_prompt_from_file("prompts/prompt_optimization.txt", {
        "cv_text": cv_text,
        "jd_text": jd_text
    })
    update_output(result)

def start_application():

    global output_text, audit_mode
    root = tk.Tk()
    root.title("ATS CV Optimizer Pro")
    root.geometry("1200x800")
    root.configure(bg='#f0f0f0')

    top_frame = tk.Frame(root)
    top_frame.pack(pady=10)

    model_label = tk.Label(top_frame, text="Model:")
    model_label.pack(side=tk.LEFT, padx=5)
    model_selector = ttk.Combobox(top_frame, values=["llama3:8b", "mistral:7b"], width=20)
    model_selector.set("llama3:8b")
    model_selector.pack(side=tk.LEFT)
    model_selector.bind("<<ComboboxSelected>>", lambda e: set_model(model_selector.get()))

    audit_mode = tk.StringVar(value="consultant")
    tk.Label(top_frame, text="Audit Mode:").pack(side=tk.LEFT, padx=5)
    audit_menu = ttk.Combobox(
        top_frame,
        textvariable=audit_mode,
        values=["consultant", "client", "self", "compare", "optimize"],
        width=20
    )
    audit_menu.pack(side=tk.LEFT)

    button_frame = tk.Frame(root)
    button_frame.pack(pady=10)

    tk.Button(button_frame, text="Upload CV", command=upload_cv).pack(side=tk.LEFT, padx=5)
    tk.Button(button_frame, text="Upload JD", command=upload_jd).pack(side=tk.LEFT, padx=5)
    tk.Button(button_frame, text="AI Audit", command=run_ai_audit).pack(side=tk.LEFT, padx=5)
    tk.Button(button_frame, text="Optimize CV", command=run_ai_audit).pack(side=tk.LEFT, padx=5)
    tk.Button(button_frame, text="Compare CV vs JD", command=run_scoring).pack(side=tk.LEFT, padx=5)
    tk.Button(button_frame, text="White Keywords", command=generate_white_keywords).pack(side=tk.LEFT, padx=5)
    tk.Button(button_frame, text="Cover Letter", command=generate_cover_letter).pack(side=tk.LEFT, padx=5)
    tk.Button(button_frame, text="Scoring", command=run_scoring).pack(side=tk.LEFT, padx=5)

    bottom_frame = tk.Frame(root)
    bottom_frame.pack(pady=10)

    tk.Button(bottom_frame, text="Save Result", command=save_result).pack(side=tk.LEFT, padx=5)
    tk.Button(bottom_frame, text="Export DOCX", command=export_docx).pack(side=tk.LEFT, padx=5)
    tk.Button(bottom_frame, text="Clear All Fields", command=clear_fields).pack(side=tk.LEFT, padx=5)
    tk.Button(bottom_frame, text="License Check", command=check_license).pack(side=tk.LEFT, padx=5)

    output_text = scrolledtext.ScrolledText(root, height=30, width=140)
    output_text.pack(padx=10, pady=10)

    root.mainloop()

if __name__ == "__main__":
    start_application()